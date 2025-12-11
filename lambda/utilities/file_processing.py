#   Copyright Amazon.com, Inc. or its affiliates. All Rights Reserved.
#
#   Licensed under the Apache License, Version 2.0 (the "License").
#   You may not use this file except in compliance with the License.
#   You may obtain a copy of the License at
#
#       http://www.apache.org/licenses/LICENSE-2.0
#
#   Unless required by applicable law or agreed to in writing, software
#   distributed under the License is distributed on an "AS IS" BASIS,
#   WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
#   See the License for the specific language governing permissions and
#   limitations under the License.

"""Helper functions to parse documents for ingestion into RAG vector store."""
import logging
import os
from io import BytesIO
from urllib.parse import urlparse

import boto3
import docx
from botocore.exceptions import ClientError
from langchain_core.documents import Document
from models.domain_objects import IngestionJob
from pypdf import PdfReader
from pypdf.errors import PdfReadError
from utilities.chunking_strategy_factory import ChunkingStrategyFactory
from utilities.constants import DOCX_FILE, PDF_FILE, RICH_TEXT_FILE, TEXT_FILE
from utilities.exceptions import RagUploadException

logger = logging.getLogger(__name__)
session = boto3.Session()
s3 = session.client("s3", region_name=os.environ["AWS_REGION"])


def _get_metadata(s3_uri: str, name: str, metadata: dict | None = None) -> dict:
    """
    Create metadata dictionary for a document.

    Args:
        s3_uri: S3 URI of the document
        name: Name of the document
        metadata: Optional additional metadata to merge into the result

    Returns:
        Dictionary containing document metadata
    """
    base_metadata = {"source": s3_uri, "name": name}

    # Merge additional metadata if provided
    if metadata:
        base_metadata.update(metadata)

    return base_metadata


def _get_s3_uri(bucket: str, key: str) -> str:
    return f"s3://{bucket}/{key}"


def _extract_text_by_content_type(content_type: str, s3_object: dict) -> str:
    extraction_functions = {
        PDF_FILE: _extract_pdf_content,
        DOCX_FILE: _extract_docx_content,
        TEXT_FILE: _extract_text_content,
        RICH_TEXT_FILE: _extract_text_content,
    }

    extraction_function = extraction_functions.get(content_type)
    if extraction_function:
        return extraction_function(s3_object)
    else:
        logger.error(f"File has unsupported content type: {content_type}")
        raise RagUploadException("Unsupported file type")


def _extract_pdf_content(s3_object: dict) -> str:
    """Return text extracted from PDF.

    Extracts text content from a PDF file in an S3 object.
    Cleans up common PDF extraction artifacts to reduce payload size.

    Parameters
    ----------
    s3_object (dict): an S3 object containing a PDF file body

    Returns
    -------
    str: The extracted text from the PDF file.
    """
    file_content = s3_object["Body"].read()
    pdf_file = BytesIO(file_content)

    try:
        pdf_reader = PdfReader(pdf_file)
    except PdfReadError as e:
        logger.error(f"Error reading PDF file: {e}")
        raise

    # Extract text from all pages
    extracted_text = "".join(page.extract_text() or "" for page in pdf_reader.pages)
    
    # Clean up common PDF extraction artifacts to reduce payload size
    cleaned_text = _clean_pdf_text(extracted_text)
    
    logger.info(f"PDF text extraction: original length={len(extracted_text)}, cleaned length={len(cleaned_text)}")
    return cleaned_text


def _clean_pdf_text(text: str) -> str:
    """
    Clean up common PDF extraction artifacts to reduce payload size and improve chunking.
    
    Args:
        text: Raw text extracted from PDF
        
    Returns:
        Cleaned text with reduced artifacts
    """
    import re
    
    # Remove excessive whitespace and normalize line breaks
    text = re.sub(r'\s+', ' ', text)  # Replace multiple whitespace with single space
    text = re.sub(r'\n\s*\n', '\n\n', text)  # Normalize paragraph breaks
    
    # Remove common PDF artifacts
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\xff]', '', text)  # Remove control characters
    text = re.sub(r'[^\x00-\x7F]+', ' ', text)  # Replace non-ASCII with space (optional - be careful with international text)
    
    # Remove excessive punctuation that might come from table borders or formatting
    text = re.sub(r'[_\-=]{3,}', '', text)  # Remove lines of underscores, dashes, equals
    text = re.sub(r'\.{3,}', '...', text)  # Normalize excessive dots
    
    # Clean up spacing around punctuation
    text = re.sub(r'\s+([,.!?;:])', r'\1', text)  # Remove space before punctuation
    text = re.sub(r'([,.!?;:])\s+', r'\1 ', text)  # Normalize space after punctuation
    
    # Remove empty lines and trim
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    
    return '\n'.join(lines)


def _extract_docx_content(s3_object: dict) -> str:
    """Return text extracted from docx document.

    Extracts text content from a docx file in an S3 object.

    Parameters
    ----------
    s3_object (dict): an S3 object containing a docx file body

    Returns
    -------
    str: The extracted text from the docx file.
    """
    streaming_body = s3_object["Body"]
    file_as_bytes = streaming_body.read()
    bytes_as_file_like = BytesIO(file_as_bytes)

    doc = docx.Document(docx=bytes_as_file_like)

    output = "\n".join(para.text for para in doc.paragraphs)
    return output


def _extract_text_content(s3_object: dict) -> str:
    """
    Extracts text content from an S3 object. Decode as
    utf-8 to properly read special characters

    Parameters
    ----------
    s3_object (dict): an S3 object containing a text file body.
    """
    return s3_object["Body"].read().decode("utf-8", errors="replace")


def generate_chunks(ingestion_job: IngestionJob) -> list[Document]:
    """Generate chunks from an ingestion job using the configured chunking strategy.

    Parameters
    ----------
    ingestion_job : IngestionJob
        Ingestion job containing file information and chunking strategy

    Returns
    -------
    list[Document]
        List of document chunks for the processed file

    Raises
    ------
    RagUploadException
        If S3 path is invalid or file processing fails
    ValueError
        If chunking strategy is not supported
    """
    # Parse S3 URI using urlparse
    parsed_uri = urlparse(ingestion_job.s3_path)
    if not parsed_uri.netloc or not parsed_uri.path:
        logger.error(f"Invalid S3 path format: {ingestion_job.s3_path}")
        raise RagUploadException("Invalid S3 path format")

    bucket = parsed_uri.netloc
    key = parsed_uri.path.lstrip("/")

    content_type = key.split(".")[-1].lower()
    try:
        s3_object = s3.get_object(Bucket=bucket, Key=key)
    except ClientError as e:
        logger.error(f"Error getting object from S3: {key}")
        raise e

    # Extract text and create initial document
    extracted_text = _extract_text_by_content_type(content_type=content_type, s3_object=s3_object)
    basename = os.path.basename(ingestion_job.s3_path)

    # Pass metadata from IngestionJob to be merged into document metadata
    docs = [
        Document(
            page_content=extracted_text,
            metadata=_get_metadata(s3_uri=ingestion_job.s3_path, name=basename, metadata=ingestion_job.metadata),
        )
    ]

    # Optimize chunking strategy for PDFs if using small chunks
    optimized_strategy = _optimize_chunking_for_file_type(ingestion_job.chunk_strategy, content_type, len(extracted_text))

    # Use factory to chunk documents based on strategy
    logger.info(f"Processing document with chunking strategy: {optimized_strategy.type}")
    doc_chunks = ChunkingStrategyFactory.chunk_documents(docs, optimized_strategy)

    # Update part number of doc metadata
    for i, doc in enumerate(doc_chunks):
        doc.metadata["part"] = i + 1

    logger.info(f"Generated {len(doc_chunks)} chunks for document: {basename} (text length: {len(extracted_text)})")
    return doc_chunks


def _optimize_chunking_for_file_type(strategy, content_type: str, text_length: int):
    """
    Optimize chunking strategy based on file type and content length to reduce payload issues.
    
    Args:
        strategy: Original chunking strategy
        content_type: File extension/type
        text_length: Length of extracted text
        
    Returns:
        Optimized chunking strategy
    """
    from models.domain_objects import FixedChunkingStrategy
    
    logger.info(f"Chunk optimization analysis - File type: {content_type}, Text length: {text_length:,} chars, "
               f"Original strategy: {strategy.type if hasattr(strategy, 'type') else 'unknown'}")
    
    if hasattr(strategy, 'size'):
        logger.info(f"Original chunk config - Size: {strategy.size}, Overlap: {strategy.overlap}")
        original_estimated_chunks = text_length // (strategy.size - strategy.overlap) if strategy.size > strategy.overlap else text_length // strategy.size
        logger.info(f"Original estimated chunks: {original_estimated_chunks}")
    
    # For PDFs with small chunk sizes, increase chunk size to reduce number of chunks
    if content_type == "pdf" and hasattr(strategy, 'size') and strategy.size < 800:
        # Calculate optimal chunk size to keep number of chunks reasonable
        target_chunks = min(200, max(10, text_length // 2000))  # Target 10-200 chunks
        optimal_size = max(800, text_length // target_chunks)
        optimal_size = min(optimal_size, 2000)  # Cap at 2000 chars
        
        # Keep overlap proportional but reasonable
        optimal_overlap = min(strategy.overlap, optimal_size // 10)
        
        # Calculate estimated chunks with new strategy
        estimated_chunks = text_length // (optimal_size - optimal_overlap) if optimal_size > optimal_overlap else text_length // optimal_size
        
        logger.info(f"PDF OPTIMIZATION APPLIED:")
        logger.info(f"  Original: size={strategy.size}, overlap={strategy.overlap}")
        logger.info(f"  Optimized: size={optimal_size}, overlap={optimal_overlap}")
        logger.info(f"  Estimated chunks: {original_estimated_chunks} → {estimated_chunks}")
        logger.info(f"  Chunk reduction: {((original_estimated_chunks - estimated_chunks) / original_estimated_chunks * 100):.1f}%")
        
        return FixedChunkingStrategy(size=optimal_size, overlap=optimal_overlap)
    else:
        logger.info(f"No optimization applied - File type: {content_type}, Chunk size: {getattr(strategy, 'size', 'N/A')}")
    
    return strategy
